"""Anonymisation de fichiers .docx en préservant la mise en forme Word."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Callable, Iterable, Iterator

try:
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
except ImportError:  # pragma: no cover - optional at install time
    Document = None  # type: ignore[misc, assignment]
    Table = None  # type: ignore[misc, assignment]
    Paragraph = None  # type: ignore[misc, assignment]


@dataclass(frozen=True)
class DocxRedactStats:
    paragraphs_processed: int
    paragraphs_changed: int


def docx_available() -> bool:
    return Document is not None


def is_docx_path(path: Path | str) -> bool:
    return Path(path).suffix.lower() == ".docx"


def _require_docx() -> None:
    if Document is None:
        raise RuntimeError(
            "Le support Word nécessite python-docx. "
            "Réinstallez l’application ou contactez le support."
        )


def _iter_table_paragraphs(table: Table) -> Iterator[Paragraph]:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested in cell.tables:
                yield from _iter_table_paragraphs(nested)


def _iter_container_paragraphs(container) -> Iterator[Paragraph]:
    """Paragraphes et tableaux d’un corps de document, en-tête ou pied de page."""
    for paragraph in container.paragraphs:
        yield paragraph
    for table in container.tables:
        yield from _iter_table_paragraphs(table)


def iter_all_paragraphs(doc: Document) -> Iterator[Paragraph]:
    """Tous les paragraphes du corps, tableaux, en-têtes et pieds de page."""
    yield from _iter_container_paragraphs(doc)
    for section in doc.sections:
        for part in (section.header, section.footer):
            if part is None:
                continue
            try:
                if not part.is_linked_to_previous:
                    yield from _iter_container_paragraphs(part)
            except AttributeError:
                yield from _iter_container_paragraphs(part)


def replace_paragraph_text(paragraph: Paragraph, new_text: str) -> None:
    """
    Remplace le texte d’un paragraphe en conservant le style du premier segment
    (gras, italique, police, etc.). Les styles de paragraphe (titre, liste) sont inchangés.
    """
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def extract_plaintext(path: Path, *, max_chars: int | None = 50_000) -> str:
    """Extrait le texte pour l’aperçu dans l’interface."""
    _require_docx()
    doc = Document(str(path))
    parts: list[str] = []
    total = 0
    for paragraph in iter_all_paragraphs(doc):
        text = paragraph.text
        if not text.strip():
            continue
        if max_chars is not None and total + len(text) + 2 > max_chars:
            parts.append("… (aperçu tronqué)")
            break
        parts.append(text)
        total += len(text) + 2
    return "\n\n".join(parts)


def default_output_path(input_path: Path) -> Path:
    return input_path.with_name(f"{input_path.stem}_anonymise.docx")


def redact_docx_file(
    input_path: Path,
    output_path: Path,
    redact_fn: Callable[[str], str],
) -> DocxRedactStats:
    """
    Copie anonymisée du .docx vers output_path.

    Chaque paragraphe est traité séparément : la mise en forme Word (styles,
    tableaux, en-têtes) est conservée ; le gras/italique à l’intérieur d’un
    paragraphe suit le premier segment de texte.
    """
    _require_docx()
    doc = Document(str(input_path))
    processed = 0
    changed = 0
    for paragraph in iter_all_paragraphs(doc):
        original = paragraph.text
        if not original.strip():
            continue
        processed += 1
        redacted = redact_fn(original)
        if redacted != original:
            replace_paragraph_text(paragraph, redacted)
            changed += 1
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return DocxRedactStats(
        paragraphs_processed=processed,
        paragraphs_changed=changed,
    )
