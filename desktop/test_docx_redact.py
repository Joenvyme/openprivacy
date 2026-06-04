"""Tests unitaires pour docx_redact (sans charger le modèle OPF)."""

from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

pytest = None  # unittest only

try:
    from docx import Document

    try:
        from .docx_redact import (
            default_output_path,
            extract_plaintext,
            redact_docx_file,
            replace_paragraph_text,
        )
    except ImportError:
        from docx_redact import (
            default_output_path,
            extract_plaintext,
            redact_docx_file,
            replace_paragraph_text,
        )
except ImportError:
    Document = None


@unittest.skipIf(Document is None, "python-docx non installé")
class DocxRedactTests(unittest.TestCase):
    def test_replace_paragraph_preserves_run_style(self) -> None:
        doc = Document()
        p = doc.add_paragraph()
        run_bold = p.add_run("Jean ")
        run_bold.bold = True
        p.add_run("Dupont")
        replace_paragraph_text(p, "<PRIVATE_PERSON>")
        self.assertEqual(p.text, "<PRIVATE_PERSON>")
        self.assertTrue(p.runs[0].bold)

    def test_redact_docx_roundtrip(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            src = Path(tmp) / "test.docx"
            out = default_output_path(src)
            doc = Document()
            doc.add_paragraph("Contact : jean.dupont@example.com")
            doc.save(str(src))

            def fake_redact(text: str) -> str:
                return text.replace("jean.dupont@example.com", "<PRIVATE_EMAIL>")

            stats = redact_docx_file(src, out, fake_redact)
            self.assertEqual(stats.paragraphs_processed, 1)
            self.assertEqual(stats.paragraphs_changed, 1)
            plain = extract_plaintext(out)
            self.assertIn("<PRIVATE_EMAIL>", plain)
            self.assertNotIn("jean.dupont@example.com", plain)


if __name__ == "__main__":
    unittest.main()
